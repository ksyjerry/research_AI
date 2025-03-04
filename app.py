import streamlit as st
import os
from openai import OpenAI
from step1_feedback.feedback import generate_feedback
from step2_research.research import deep_research
from step3_reporting.reporting import write_final_report
from PIL import Image
import time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# 페이지 설정
st.set_page_config(
    page_title="Research AI",
    page_icon="🔍",
    layout="wide"
)

# CSS 스타일 추가
st.markdown("""
    <style>
    .logo-title-container {
        display: flex;
        align-items: center;
        gap: 10px;
        margin-bottom: 20px;
    }
    .logo-title-container h1 {
        margin: 0;
        padding: 0;
    }
    .sidebar .stMarkdown {
        font-size: 0.9em;
    }
    .sidebar .stMarkdown h3 {
        font-size: 1.1em;
        margin-bottom: 0.5em;
    }
    .sidebar .stMarkdown ul {
        margin-top: 0.3em;
        margin-bottom: 0.3em;
    }
    .sidebar .stMarkdown li {
        margin-top: 0.2em;
        margin-bottom: 0.2em;
    }
    /* 표 스타일 개선 */
    .stMarkdown table {
        width: 100%;
        margin: 1em 0;
        border-collapse: collapse;
    }
    .stMarkdown th {
        background-color: #f0f2f6;
        padding: 0.5em;
        border: 1px solid #ddd;
        font-weight: bold;
    }
    .stMarkdown td {
        padding: 0.5em;
        border: 1px solid #ddd;
        vertical-align: top;
    }
    .stMarkdown tr:nth-child(even) {
        background-color: #f8f9fa;
    }
    .stMarkdown tr:hover {
        background-color: #f0f2f6;
    }
    </style>
""", unsafe_allow_html=True)

# 세션 상태 초기화
if 'current_question' not in st.session_state:
    st.session_state.current_question = 0
if 'answers' not in st.session_state:
    st.session_state.answers = []
if 'feedback_questions' not in st.session_state:
    st.session_state.feedback_questions = None
if 'combined_query' not in st.session_state:
    st.session_state.combined_query = None
if 'research_completed' not in st.session_state:
    st.session_state.research_completed = False
if 'research_results' not in st.session_state:
    st.session_state.research_results = None
if 'report' not in st.session_state:
    st.session_state.report = None
if 'current_learning' not in st.session_state:
    st.session_state.current_learning = 0
if 'report_sections' not in st.session_state:
    st.session_state.report_sections = []
if 'research_started' not in st.session_state:
    st.session_state.research_started = False
if 'research_progress' not in st.session_state:
    st.session_state.research_progress = []

# 환경 변수 설정
os.environ["OPENAI_API_KEY"] = st.secrets["openai"]["OPENAI_API_KEY"]
os.environ["FIRECRAWL_API_KEY"] = st.secrets["firecrawl"]["FIRECRAWL_API_KEY"]

def create_word_document(report_text, query):
    doc = Document()
    
    # 제목 스타일 설정
    title = doc.add_heading('Research Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 초기 질문 추가
    doc.add_heading('초기 연구 질문', level=1)
    doc.add_paragraph(query)
    
    # 보고서 내용 추가
    doc.add_heading('연구 결과', level=1)
    
    # 마크다운 표 패턴 매칭을 위한 정규식
    table_pattern = r'\|([^\n]+)\|\n\|[-|]+\|\n(((?!\n\n|\n#).*\|\n)*)'
    url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
    
    current_text = ""
    sections = report_text.split('\n\n')
    
    for section in sections:
        if section.strip():
            # 섹션이 표를 포함하는지 확인
            if '|' in section and '-|-' in section:
                # 현재까지의 텍스트 추가
                if current_text.strip():
                    p = doc.add_paragraph(current_text.strip())
                    p.style.font.name = '맑은 고딕'
                    p.style.font.size = Pt(11)
                    current_text = ""
                
                # 표 처리
                rows = [row.strip() for row in section.split('\n') if row.strip() and '|' in row]
                if len(rows) >= 2:  # 헤더와 구분선, 데이터가 있는 경우
                    # 헤더 처리
                    headers = [cell.strip() for cell in rows[0].split('|') if cell.strip()]
                    
                    # 표 생성
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # 헤더 채우기
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        header_cells[i].paragraphs[0].runs[0].font.bold = True
                        header_cells[i].paragraphs[0].runs[0].font.name = '맑은 고딕'
                    
                    # 데이터 행 추가
                    for row in rows[2:]:  # 구분선 다음부터
                        cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                        if cells:
                            new_row = table.add_row().cells
                            for i, cell in enumerate(cells):
                                if i < len(new_row):
                                    new_row[i].text = cell
                                    new_row[i].paragraphs[0].runs[0].font.name = '맑은 고딕'
                    
                    # 표 다음에 빈 줄 추가
                    doc.add_paragraph()
            
            # 섹션 제목인 경우
            elif section.startswith('#'):
                # 현재까지의 텍스트 추가
                if current_text.strip():
                    p = doc.add_paragraph(current_text.strip())
                    p.style.font.name = '맑은 고딕'
                    p.style.font.size = Pt(11)
                    current_text = ""
                
                level = section.count('#')
                text = section.lstrip('#').strip()
                heading = doc.add_heading(text, level=level)
                for run in heading.runs:
                    run.font.name = '맑은 고딕'
            
            else:
                # URL을 하이퍼링크로 변환
                text = section
                urls = re.findall(url_pattern, text)
                
                if urls:
                    # 현재까지의 텍스트 추가
                    if current_text.strip():
                        p = doc.add_paragraph(current_text.strip())
                        p.style.font.name = '맑은 고딕'
                        p.style.font.size = Pt(11)
                        current_text = ""
                    
                    p = doc.add_paragraph()
                    start = 0
                    for url in urls:
                        # URL 이전의 텍스트 추가
                        pos = text.find(url, start)
                        if pos > start:
                            p.add_run(text[start:pos]).font.name = '맑은 고딕'
                        # URL을 하이퍼링크로 추가
                        p.add_hyperlink(url, url, '0563C1')
                        start = pos + len(url)
                    # 남은 텍스트 추가
                    if start < len(text):
                        p.add_run(text[start:]).font.name = '맑은 고딕'
                else:
                    current_text += text + "\n"
    
    # 마지막 텍스트 추가
    if current_text.strip():
        p = doc.add_paragraph(current_text.strip())
        p.style.font.name = '맑은 고딕'
        p.style.font.size = Pt(11)
    
    # 문서를 바이트 스트림으로 변환
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def sanitize_filename(filename):
    # 파일명에서 사용할 수 없는 문자 제거
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    # 파일명이 너무 길 경우 자동으로 줄임
    if len(filename) > 50:
        filename = filename[:50]
    return filename

# 로고와 제목 표시
st.markdown('<div class="logo-title-container">', unsafe_allow_html=True)
col1, col2 = st.columns([1, 8])
with col1:
    st.image("PwC.jpg", width=80)
with col2:
    st.title("Research AI")
st.markdown('</div>', unsafe_allow_html=True)

# OpenAI 클라이언트 초기화
client = OpenAI()

# 모델 설정
feedback_model = "gpt-4o-mini"    
research_model = "o3-mini"
reporting_model = "o3-mini"

# 사이드바에 설정 추가
with st.sidebar:
    st.header("설정")
    
    st.markdown("""
    ### 연구 범위 (1-5)
    - 연구 주제와 관련된 하위 주제의 수를 의미합니다
    - 1: 가장 핵심적인 주제만 연구
    - 5: 주제와 관련된 모든 하위 주제를 포괄적으로 연구
    - 기본값: 2
    """)
    breadth = st.number_input("연구 범위", min_value=1, max_value=5, value=2)
    
    st.markdown("""
    ### 연구 깊이 (1-5)
    - 각 하위 주제에 대한 연구의 상세도를 의미합니다
    - 1: 기본적인 정보만 수집
    - 5: 매우 상세한 분석과 심층적인 정보 수집
    - 기본값: 2
    """)
    depth = st.number_input("연구 깊이", min_value=1, max_value=5, value=2)
    
    st.markdown("""
    ---
    **참고사항**
    - 범위와 깊이를 높일수록 연구 시간이 길어집니다
    - 처음에는 기본값(2)으로 시작하는 것을 추천합니다
    """)

# 메인 컨텐츠
st.header("연구 주제 입력")
query = st.text_input("어떤 주제에 대해 리서치하시겠습니까?")

if query:
    # 1단계: 추가 질문 생성
    st.header("1단계: 추가 질문")
    
    # 질문이 아직 생성되지 않은 경우에만 생성
    if st.session_state.feedback_questions is None:
        st.session_state.feedback_questions = generate_feedback(query, client, feedback_model, max_feedbacks=3)
    
    if st.session_state.feedback_questions:
        st.write("다음 질문에 답변해 주세요:")
        
        # 현재 질문까지의 답변 표시
        for i in range(st.session_state.current_question):
            st.text_input(f"질문 {i+1}: {st.session_state.feedback_questions[i]}", 
                         value=st.session_state.answers[i], disabled=True)
        
        # 현재 질문 입력
        current_answer = st.text_input(f"질문 {st.session_state.current_question + 1}: {st.session_state.feedback_questions[st.session_state.current_question]}")
        
        # 답변 제출 버튼
        if st.button("답변 제출"):
            if current_answer:
                st.session_state.answers.append(current_answer)
                if st.session_state.current_question < len(st.session_state.feedback_questions) - 1:
                    st.session_state.current_question += 1
                    st.rerun()
                else:
                    # 모든 답변이 완료된 경우
                    st.session_state.combined_query = f"초기 질문: {query}\n"
                    for i in range(len(st.session_state.feedback_questions)):
                        st.session_state.combined_query += f"\n{i+1}. 질문: {st.session_state.feedback_questions[i]}\n"
                        st.session_state.combined_query += f"   답변: {st.session_state.answers[i]}\n"
                    st.rerun()
    else:
        st.warning("추가 질문이 생성되지 않았습니다.")

# 최종 질문과 연구 시작 버튼 표시
if st.session_state.combined_query:
    if st.button("연구 시작"):
        st.session_state.research_started = True
        st.session_state.research_completed = False
        st.session_state.current_learning = 0
        st.session_state.report_sections = []
        st.session_state.research_results = None
        st.session_state.research_progress = []
        st.rerun()

# 연구 진행 상태 표시
if st.session_state.research_completed:
    with st.expander("연구 진행 과정", expanded=True):
        st.header("2단계: 딥리서치")
        st.write("연구 결과:")
        for learning in st.session_state.research_results["learnings"]:
            st.write(f"• {learning}")
        
        st.header("3단계: 최종 보고서 작성")
        st.markdown(st.session_state.report)
        
        # 보고서 저장 및 다운로드
        os.makedirs("output", exist_ok=True)
        
        # Word 문서 다운로드
        doc_io = create_word_document(st.session_state.report, query)
        safe_filename = sanitize_filename(f"리서치보고서_{query}")
        st.download_button(
            label="보고서 다운로드",
            data=doc_io,
            file_name=f"{safe_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
elif st.session_state.combined_query and st.session_state.research_started and not st.session_state.research_completed:
    # 2단계: 심층 연구
    st.header("2단계: 딥리서치")
    research_placeholder = st.empty()
    
    if st.session_state.research_results is None:
        # 진행 상황을 표시할 컨테이너 생성
        progress_container = st.container()
        
        with progress_container:
            with st.spinner("🔍 연구를 진행하는 중입니다..."):
                progress_placeholder = st.empty()
                
                def update_progress(message):
                    st.session_state.research_progress.append(message)
                    progress_text = "\n".join(st.session_state.research_progress)
                    progress_placeholder.markdown(f"```\n{progress_text}\n```")
                
                # deep_research 함수 호출 전에 콜백 함수 설정
                def research_callback(message):
                    update_progress(message)
                
                st.session_state.research_results = deep_research(
                    query=st.session_state.combined_query,
                    breadth=breadth,
                    depth=depth,
                    client=client,
                    model=research_model,
                    callback=research_callback
                )
    
    # 연구 진행 상황 스트리밍 표시
    with st.expander("연구 진행 상황", expanded=True):
        st.subheader("1. 초기 연구 질문 분석")
        st.write(st.session_state.combined_query)
        
        if st.session_state.research_results:
            st.subheader("2. 검색 키워드 생성")
            for i, serp_query in enumerate(st.session_state.research_results.get("serp_queries", []), 1):
                st.write(f"키워드 {i}: {serp_query.query}")
                st.write(f"목적: {serp_query.research_goal}")
                time.sleep(0.3)
            
            st.subheader("3. 연구 결과")
            for i in range(len(st.session_state.research_results["learnings"])):
                if i >= st.session_state.current_learning:
                    st.write(f"• {st.session_state.research_results['learnings'][i]}")
                    st.session_state.current_learning = i + 1
                    time.sleep(0.5)
            
            if st.session_state.research_results.get("visited_urls"):
                st.subheader("4. 참조 URL")
                for url in st.session_state.research_results["visited_urls"]:
                    st.write(f"• {url}")
                    time.sleep(0.2)

    # 3단계: 최종 보고서
    if st.session_state.current_learning == len(st.session_state.research_results["learnings"]):
        st.header("3단계: 보고서 작성")
        report_placeholder = st.empty()
        
        if not st.session_state.report:
            with st.spinner("📝 보고서를 작성하는 중입니다..."):
                st.session_state.report = write_final_report(
                    prompt=st.session_state.combined_query,
                    learnings=st.session_state.research_results["learnings"],
                    visited_urls=st.session_state.research_results["visited_urls"],
                    client=client,
                    model=reporting_model
                )
                
                # 보고서를 섹션별로 나누어 스트리밍
                sections = st.session_state.report.split('\n\n')
                for section in sections:
                    if section.strip():
                        st.markdown(section)
                        time.sleep(0.3)
            
            st.session_state.research_completed = True
            st.rerun() 